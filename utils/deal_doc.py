#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# @File  : deal_doc.py
# @Author: itnoobzzy
# @Date  : 2021/3/20
# @Desc  : 处理word文档操作

from docx import Document
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor, Inches, Pt, Length, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

class DocxCaseInfo:
    def __init__(self, path):
        self.path = path
        self.document = Document()
        self.document.styles['Normal'].font.name = u'宋体'
        self.document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        self.table = ''
        self.firstflag = True

    # 添加表格
    def add_table(self, tablename):
        if self.firstflag is True:
            self.firstflag = False
        else:
            self.document.add_page_break()
        self.create_head(tablename)
        # 创建表格
        self.table = self.document.add_table(rows=12, cols=4, style='Table Grid')
        self.init_table()

    def add_table_1(self, tablename):
        if self.firstflag is True:
            self.firstflag = False
        else:
            self.document.add_page_break()
        self.create_head(tablename)
        # 创建表格
        self.table = self.document.add_table(rows=2, cols=5, style='Table Grid')
        self.init_table_1()

    def add_table_2(self, tablename, rows):
        if self.firstflag is True:
            self.firstflag = False
        else:
            self.document.add_page_break()
        self.create_head(tablename)
        # 创建表格
        self.table = self.document.add_table(rows=rows, cols=6, style='Table Grid')
        self.init_table_2()

    def add_table_3(self, tablename):
        if self.firstflag is True:
            self.firstflag = False
        else:
            self.document.add_page_break()
        self.create_head(tablename)

    # 初始化表上文字
    def create_head(self,tablename):

        p = self.document.add_paragraph('', style='Heading1')
        r = p.add_run(tablename)
        r.font.name = "宋体"


        r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        # 上下间距
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        # 首行缩进
        p.paragraph_format.first_line_indent = Inches(0.3)
        r.font.color.rgb = RGBColor(0, 0, 0)

    # 初始化表格
    def init_table(self):
        for rowcnt in range(len(self.table.rows)):
            row = self.table.rows[rowcnt]
            row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            self._row_deal(row, rowcnt)

            if rowcnt == 0:
                row.height = Cm(0.7)
                r = row.cells[0].paragraphs[0].add_run('需求编号')
                r.bold = True
                r = row.cells[2].paragraphs[0].add_run('编制人')
                r.bold = True
            elif rowcnt == 1:
                row.height = Cm(0.7)
                r = row.cells[0].paragraphs[0].add_run('版本')
                r.bold = True
                r = row.cells[2].paragraphs[0].add_run('审定人')
                r.bold = True
            elif rowcnt == 2:
                row.height = Cm(0.7)
                r = row.cells[0].paragraphs[0].add_run('测试用例编号')
                r.bold = True
                r = row.cells[2].paragraphs[0].add_run('测试名称')
                r.bold = True
            elif rowcnt == 3:
                row.height = Cm(1.5)
                r = row.cells[0].paragraphs[0].add_run('测试目的')
                r.bold = True
            elif rowcnt == 4:
                row.height = Cm(2)
                r = row.cells[0].paragraphs[0].add_run('测试条件说明')
                r.bold = True
            elif rowcnt == 5:
                row.height = Cm(2)
                r = row.cells[0].paragraphs[0].add_run('测试步骤')
                r.bold = True
            elif rowcnt == 6:
                row.height = Cm(2)
                r = row.cells[0].paragraphs[0].add_run('期待输出结果')
                r.bold = True
            elif rowcnt == 7:
                row.height = Cm(2)
                r = row.cells[0].paragraphs[0].add_run('实际输出结果')
                r.bold = True
            elif rowcnt == 8:
                row.height = Cm(2)
                r = row.cells[0].paragraphs[0].add_run('测试报告结论')
                r.bold = True
            elif rowcnt == 9:
                row.height = Cm(0.7)
                r = row.cells[0].paragraphs[0].add_run('测试人员')
                r.bold = True
                r = row.cells[2].paragraphs[0].add_run('测试日期')
                r.bold = True
            elif rowcnt == 10:
                row.height = Cm(0.7)
                r = row.cells[0].paragraphs[0].add_run('审查人员')
                r.bold = True
                r = row.cells[2].paragraphs[0].add_run('审查时间')
                r.bold = True
            elif rowcnt == 11:
                row.height = Cm(2)
                r = row.cells[0].paragraphs[0].add_run('审查结果')
                r.bold = True

    def _row_deal(self, row, rowcnt):
        if 2 < rowcnt <= 8 or rowcnt > 10:
            # 3行后需要合并
            row.cells[1].merge(row.cells[3])
        else:
            row.cells[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def init_table_1(self):
        # print(len(self.table.rows))
        for rowcnt in range(len(self.table.rows)):
            # print(rowcnt)
            row = self.table.rows[rowcnt]

            if rowcnt == 0:
                row.height = Cm(0.7)
                r = row.cells[0].paragraphs[0].add_run('测试用例总数')
                r.bold = True
                r = row.cells[1].paragraphs[0].add_run('测试通过用例数')
                r.bold = True
                r = row.cells[2].paragraphs[0].add_run('测试不通过用例数')
                r.bold = True
                r = row.cells[3].paragraphs[0].add_run('测试时间')
                r.bold = True
                r = row.cells[4].paragraphs[0].add_run('测试人员')
                r.bold = True

    def init_table_2(self):
        # print(len(self.table.rows))
        for rowcnt in range(len(self.table.rows)):
            # print(rowcnt)
            row = self.table.rows[rowcnt]

            if rowcnt == 0:
                row.height = Cm(0.7)
                r = row.cells[0].paragraphs[0].add_run('序号')
                r.bold = True
                r = row.cells[1].paragraphs[0].add_run('测试编号')
                r.bold = True
                r = row.cells[2].paragraphs[0].add_run('测试名称')
                r.bold = True
                r = row.cells[3].paragraphs[0].add_run('是否通过')
                r.bold = True
                r = row.cells[4].paragraphs[0].add_run('测试时间')
                r.bold = True
                r = row.cells[5].paragraphs[0].add_run('测试人员')
                r.bold = True

    """
    针对具体单个表格添加具体的值
    row 行号-1
    cell 列号-1
    value 填充的值
    bold 是否加粗，默认False为不加粗
    """
    def write_table(self, row, cell, value, bold=False):
        r = self.table.rows[row].cells[cell].paragraphs[0].add_run(value)
        r.bold = bold

    def save(self):

        self.document.save(self.path)
        return self.path